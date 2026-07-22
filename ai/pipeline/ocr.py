"""
OpsPilot — OCR Engine
========================
Extracts raw text from PDFs, Images, Word docs, Excel files, and plain text.
"""

import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)

class OCREngine:
    def __init__(self):
        # Initialize Tesseract, PyMuPDF, etc. if available.
        pass

    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from the given file based on its extension.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            return self._extract_from_image(file_path)
        elif ext in ['.txt', '.md', '.csv']:
            return self._extract_from_text(file_path)
        elif ext == '.docx':
            return self._extract_from_docx(file_path)
        elif ext == '.doc':
            return self._extract_from_doc(file_path)
        elif ext in ['.xls', '.xlsx']:
            return self._extract_from_excel(file_path)
        else:
            logger.warning(f"Unsupported file type for OCR: {ext}")
            return f"[Unsupported file type: {ext}]"

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF (fitz) or pdfplumber."""
        text = ""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text("text") + "\n"
            doc.close()
        except ImportError:
            logger.warning("PyMuPDF (fitz) not installed. Falling back to stub extraction.")
            text = f"[PDF Text Extraction Stub for {os.path.basename(file_path)}]\n"
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise

        return text.strip()

    def _extract_from_image(self, file_path: str) -> str:
        """Extract text from images using Gemini Vision for P&ID intelligence."""
        import base64
        from langchain_core.messages import HumanMessage
        from ai.llm_factory import get_llm

        try:
            with open(file_path, "rb") as image_file:
                encoded_string = base64.b64encode(image_file.read()).decode("utf-8")

            ext = os.path.splitext(file_path)[1].lower().replace('.', '')
            if ext == 'jpg':
                ext = 'jpeg'

            llm = get_llm(temperature=0.1)
            message = HumanMessage(
                content=[
                    {
                        "type": "text",
                        "text": (
                            "You are an expert industrial engineer. Extract all readable text from this document. "
                            "If this is a P&ID, schematic, or diagram, carefully describe the visual components "
                            "(valves, pumps, tanks, pipes) and their connections. Output the information clearly as plain text."
                        )
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/{ext};base64,{encoded_string}"}
                    }
                ]
            )
            logger.info(f"Sending image {os.path.basename(file_path)} to Gemini Vision API")
            res = llm.invoke([message])
            return str(res.content).strip()
        except Exception as e:
            logger.error(f"Error extracting Image with Gemini: {e}")
            return f"[Image AI Extraction Failed: {str(e)}]"

    def _extract_from_text(self, file_path: str) -> str:
        """Read raw text files (TXT, CSV, MD)."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from .docx using python-docx."""
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_texts.append(row_text)
            full_text = "\n".join(paragraphs)
            if table_texts:
                full_text += "\n\n[Tables]\n" + "\n".join(table_texts)
            return full_text
        except ImportError:
            logger.warning("python-docx not installed. Run: pip install python-docx")
            return f"[DOCX Extraction Stub for {os.path.basename(file_path)}]\n"
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return f"[DOCX Extraction Failed: {str(e)}]"

    def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from legacy .doc files using antiword or textract."""
        try:
            import subprocess
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout:
                return result.stdout.strip()
            raise RuntimeError(f"antiword failed: {result.stderr}")
        except (FileNotFoundError, RuntimeError):
            # antiword not installed — try textract as fallback
            try:
                import textract
                text = textract.process(file_path)
                return text.decode('utf-8', errors='replace').strip()
            except ImportError:
                logger.warning("Neither antiword nor textract is available for .doc extraction.")
                return f"[DOC Extraction Stub for {os.path.basename(file_path)}. Install antiword or textract.]"
            except Exception as e:
                logger.error(f"Error extracting DOC with textract: {e}")
                return f"[DOC Extraction Failed: {str(e)}]"

    def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel files (.xls / .xlsx) using openpyxl or xlrd."""
        ext = os.path.splitext(file_path)[1].lower()
        text_rows = []

        try:
            if ext == '.xlsx':
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    text_rows.append(f"[Sheet: {sheet_name}]")
                    for row in ws.iter_rows(values_only=True):
                        row_text = " | ".join(str(cell) for cell in row if cell is not None)
                        if row_text.strip():
                            text_rows.append(row_text)
                wb.close()
            elif ext == '.xls':
                try:
                    import xlrd
                    wb = xlrd.open_workbook(file_path)
                    for sheet in wb.sheets():
                        text_rows.append(f"[Sheet: {sheet.name}]")
                        for row_idx in range(sheet.nrows):
                            row_text = " | ".join(str(sheet.cell_value(row_idx, col)) for col in range(sheet.ncols))
                            if row_text.strip():
                                text_rows.append(row_text)
                except ImportError:
                    # Try openpyxl as fallback for xls (may not work for older formats)
                    import openpyxl
                    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                    for sheet_name in wb.sheetnames:
                        ws = wb[sheet_name]
                        text_rows.append(f"[Sheet: {sheet_name}]")
                        for row in ws.iter_rows(values_only=True):
                            row_text = " | ".join(str(cell) for cell in row if cell is not None)
                            if row_text.strip():
                                text_rows.append(row_text)
                    wb.close()

            return "\n".join(text_rows) if text_rows else f"[Empty spreadsheet: {os.path.basename(file_path)}]"

        except ImportError as e:
            logger.warning(f"Excel library not installed ({e}). Run: pip install openpyxl xlrd")
            return f"[Excel Extraction Stub for {os.path.basename(file_path)}. Install openpyxl.]"
        except Exception as e:
            logger.error(f"Error extracting Excel file {file_path}: {e}")
            return f"[Excel Extraction Failed: {str(e)}]"
